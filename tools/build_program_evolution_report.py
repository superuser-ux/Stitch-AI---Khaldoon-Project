#!/usr/bin/env python3
from __future__ import annotations

import subprocess
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement as SharedOxmlElement
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
DOCS_DIR = ROOT / "docs"
OUT_DIR = ROOT / "output" / "doc"
TMP_DIR = ROOT / "tmp" / "docs"
LOGO = ROOT / "assets" / "brand" / "tenants" / "tanaghum" / "tanaghom-logo.png"
MD_PATH = DOCS_DIR / "14_Tanaghom_AI_Content_Department_Program_Evolution_Report_v1_0.md"
DOCX_PATH = OUT_DIR / "Tanaghom_AI_Content_Department_Program_Evolution_Report_v1_0.docx"

TITLE = "Tanaghom AI Content Department Program Evolution Report"
SUBTITLE = "Client Change Request Integration, Implementation Status & Release Roadmap\nTarget Release v1.0"
REPORT_DATE = "July 2, 2026"
VERSION = "1.2"
CLASSIFICATION = "Client Communication / Internal Governance"
SHARED_REFERENCE_BASELINE = "This issue is the current common client/internal reference set until a newer issue is formally published."

COLOR_NAVY = RGBColor(0x0B, 0x25, 0x45)
COLOR_BLUE = RGBColor(0x2E, 0x74, 0xB5)
COLOR_BLUE_DARK = RGBColor(0x1F, 0x4D, 0x78)
COLOR_MUTED = RGBColor(0x55, 0x55, 0x55)
COLOR_BORDER = "DADCE0"
COLOR_FILL = "F2F4F7"
STATUS_FILL = {
    "Implemented": "E8F3EC",
    "Partially Implemented": "FFF4DB",
    "Not Yet Implemented": "FDEBEC",
    "Unable to Verify": "F2F4F7",
    "Verified in Companion System": "EAF1FB",
    "Ready": "E8F3EC",
    "Conditional": "FFF4DB",
    "No": "FDEBEC",
}


def git_branch() -> str:
    try:
        return subprocess.check_output(
            ["git", "rev-parse", "--abbrev-ref", "HEAD"],
            cwd=ROOT,
            text=True,
        ).strip()
    except Exception:
        return "Unable to verify from repository"


def ensure_dirs() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    TMP_DIR.mkdir(parents=True, exist_ok=True)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, inches: float) -> None:
    cell.width = Inches(inches)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(int(inches * 1440)))


def set_table_fixed(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")


def set_table_borders(table, color: str = COLOR_BORDER) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def set_run_font(run, name="Calibri", size=11, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color, before, after in (
        ("Heading 1", 16, COLOR_BLUE, 16, 8),
        ("Heading 2", 13, COLOR_BLUE, 12, 6),
        ("Heading 3", 12, COLOR_BLUE_DARK, 8, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.1


def enable_update_fields_on_open(doc: Document) -> None:
    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def add_page_field(paragraph, field_code: str) -> None:
    fld_begin = SharedOxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = SharedOxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    fld_sep = SharedOxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = SharedOxmlElement("w:t")
    fld_text.text = "1"
    fld_end = SharedOxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r = paragraph.add_run()._r
    r.append(fld_begin)
    r.append(instr)
    r.append(fld_sep)
    r.append(fld_text)
    r.append(fld_end)


def add_toc_field(paragraph) -> None:
    fld_begin = SharedOxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = SharedOxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = SharedOxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = SharedOxmlElement("w:t")
    fld_text.text = (
        "1 Executive Summary; 2 Repository Maturity Assessment; 3 Current Implemented "
        "Capabilities; 4 Client Demonstration Summary; 5 Consolidated Change Requests; "
        "6 Repository Alignment Matrix; 7 Progress Dashboard; 8 Release Roadmap; "
        "9 Risks and Recommendations; 10 Client Acceptance Considerations; 11 Appendix."
    )
    fld_end = SharedOxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r = paragraph.add_run()._r
    r.append(fld_begin)
    r.append(instr)
    r.append(fld_sep)
    r.append(fld_text)
    r.append(fld_end)


def add_footer(section) -> None:
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Tanaghom Program Evolution Report | Page ")
    set_run_font(r, size=9, color=COLOR_MUTED)
    add_page_field(p, "PAGE")


def add_text(doc: Document, text: str, *, bold=False, italic=False, color=COLOR_NAVY,
             size=11, align=WD_ALIGN_PARAGRAPH.LEFT, after=6, before=0, style=None):
    p = doc.add_paragraph(style=style)
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return p


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=11, color=COLOR_NAVY)


def add_visual(doc: Document, path: Path, caption: str, *, width: float) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_after = Pt(4)
    cr = cp.add_run(caption)
    set_run_font(cr, size=9.5, color=COLOR_MUTED, italic=True)


def add_placeholder(doc: Document, title: str, note: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f"[Placeholder] {title}")
    set_run_font(r1, size=10.5, color=COLOR_NAVY, bold=True)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(6)
    r2 = p2.add_run(note)
    set_run_font(r2, size=9.25, color=COLOR_MUTED, italic=True)


def add_label_value_table(doc: Document, rows: list[tuple[str, str]], widths=(1.9, 4.6)):
    table = doc.add_table(rows=0, cols=2)
    set_table_fixed(table)
    set_table_borders(table)
    for label, value in rows:
        cells = table.add_row().cells
        set_cell_width(cells[0], widths[0])
        set_cell_width(cells[1], widths[1])
        cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(cells[0], COLOR_FILL)
        p1 = cells[0].paragraphs[0]
        p1.paragraph_format.space_after = Pt(2)
        r1 = p1.add_run(label)
        set_run_font(r1, size=10.5, color=COLOR_NAVY, bold=True)
        p2 = cells[1].paragraphs[0]
        p2.paragraph_format.space_after = Pt(2)
        r2 = p2.add_run(value)
        set_run_font(r2, size=10.5, color=COLOR_NAVY)
    doc.add_paragraph()
    return table


def add_matrix_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float],
                     status_col: int | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_fixed(table)
    set_table_borders(table)
    for idx, header in enumerate(headers):
        set_cell_width(table.rows[0].cells[idx], widths[idx])
        set_cell_shading(table.rows[0].cells[idx], COLOR_FILL)
        p = table.rows[0].cells[idx].paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(header)
        set_run_font(r, size=10, color=COLOR_NAVY, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_width(cells[idx], widths[idx])
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(value)
            set_run_font(r, size=9.5, color=COLOR_NAVY)
            if status_col is not None and idx == status_col:
                fill = STATUS_FILL.get(value, None)
                if fill:
                    set_cell_shading(cells[idx], fill)
    doc.add_paragraph()
    return table


@dataclass
class SourceRef:
    code: str
    title: str
    path: str
    use: str


BRANCH = git_branch()
ASSESSMENT_BASIS = (
    f"Current Tanaghom application build inspected on {REPORT_DATE}, including the active local "
    f"build reference `{BRANCH}`, tracked documentation, application source code, database "
    "migrations, automated verification suites, live application health checks, and selected "
    "companion-system evidence for Analytics, SDAM, and AVP supplied for this review."
)

SOURCE_REFS = [
    SourceRef("S01", "Program status and execution tracking", "README.md; BUILD_STATE.md", "Program history, current priorities, and execution continuity."),
    SourceRef("S02", "Architecture and phased roadmap", "ARCHITECTURE.md; ROADMAP.md", "Target operating model and the intended phased release path."),
    SourceRef("S03", "Continuity and demonstration notes", "HANDOFF.md", "Latest continuity note, design direction, demonstration context, and operating constraints."),
    SourceRef("S04", "Configuration baseline", "system_config.example.yaml", "Stage model, approvals, planner template, and integration declarations."),
    SourceRef("S05", "Core schema and migrations", "db/init/schema.sql; db/migrations/013-017", "Implemented data model, workflow/versioning, and approval-governance tables."),
    SourceRef("S06", "Planning implementation", "planner/plan_round.py", "28-day planning logic, ratio scaling, slot progression, and run creation behavior."),
    SourceRef("S07", "Generation implementation", "agents/run_writers.py; agents/providers.py", "Topic and script generation, revision handling, and language-quality safeguards."),
    SourceRef("S08", "Approval engine and APIs", "gates/engine.py; gates/api.py; gates/directives.py; gates/dam.py", "Approval workflow, audit model, manual lifecycle stages, and asset handoff spine."),
    SourceRef("S09", "Operational and admin UI", "dashboard/lib/review-context.tsx; dashboard/components/review/*; dashboard/components/admin/*", "Current end-user and admin user interface implementation."),
    SourceRef("S10", "Methodology loader", "loader/load_methodology.py", "Loading of methodology and content-format records into runtime and versioned tables."),
    SourceRef("S11", "Integration contracts", "integrations/contracts.py; integrations/stubs.py; docs/INTEGRATION_CONTRACTS.md", "Tanaghom-side seams for external execution and analytics connections."),
    SourceRef("S12", "Change Request 01 planning pack", "docs/10_Change_Request_01_Internal_Assessment.md; docs/11_Change_Request_01_Phased_Backlog.md; docs/12_CR01_Parallel_Execution_Map.md", "Change-request scope, backlog sequencing, and implementation framing."),
    SourceRef("S13", "Scheduling and redesign references", "docs/design/lunaris/05-content-schedule-reference.md; docs/UI_VISION.md", "Client schedule reference and the intended operating experience direction."),
    SourceRef("S14", "Quality and voice calibration references", "docs/05_Voice_Performance_Calibration.md; docs/07_Test_Drive_Real_vs_Generated.md; docs/QUALITY_BACKLOG.md", "Known content-quality expectations and open quality hardening items."),
    SourceRef("S15", "Automated verification assets", "gates/selftest.py; gates/api_selftest.py; gates/lifecycle_selftest.py; dashboard/e2e/*.ts", "Automated verification coverage used during this review."),
    SourceRef("S16", "Live UI audit evidence", "output/playwright/mobile-audit/report.json; output/playwright/mobile-audit-live/report.json; output/playwright/mobile-audit-final/report.json", "Responsive audit evidence and current UI screenshots."),
    SourceRef("S17", "Analytics companion-system review", "Brand Shield Repo: README.md; PROJECT_STATUS.md; DASHBOARD_STATUS.md; docs/SOCIAL_ANALYTICS_SCHEMA.md; docs/META_APP_SETUP.md; grafana/dashboards/*.json", "Evidence of the external analytics engine, connectors, dashboards, and operating nuance."),
    SourceRef("S18", "SDAM companion-system review", "resourcespace-stitch-stack: README.md; docker-compose.yml; docs/resourcespace-role.md; scripts/start.ps1", "Evidence of the external ResourceSpace-based smart asset-library stack."),
    SourceRef("S19", "AVP companion-system review", "Agentic-Video-Producer: README.md; docs/WORKLOG.md; apps/editor-shell; packages/mcp-server", "Evidence of the external video-production workstream and its current maturity."),
]

EXEC_SUMMARY = [
    "The Tanaghom application already demonstrates a real operating core for content planning, topic and script generation, gated human review, audit tracking, and manual handoff into production-stage workflows.",
    "Change Request 01 is materially underway inside the current build. Approval assignment foundations, reviewer-specific work visibility, workflow versioning, and methodology catalogs are present, but the full governance scope is not yet complete.",
    "The wider solution landscape is also real. Separate companion systems were verified for Analytics, Smart Digital Asset Management, and Agentic Video Production; however, those systems are not yet fully activated end-to-end inside the Tanaghom application itself.",
    "Target release v1.0 should therefore be positioned as a controlled, phased release baseline rather than a fully finished end-state. The safest sequence is Tanaghom control-plane completion first, then external-system activation and release hardening.",
    "Where evidence could not be verified directly from the inspected materials, this report states that explicitly. No implementation claim in this document has been made on assumption alone.",
]

MATURITY_ROWS = [
    ["Tanaghom core operating model", "Advanced pilot", "The current application build already supports the main editorial lifecycle from planning through gated approval and staged handoff.", "S04-S09, S15"],
    ["Governance and control features", "Medium to High", "Approval assignment foundations, workflow versioning, and methodology cataloging are present, but still require additional completion for release governance.", "S05, S08-S10, S12"],
    ["User experience and client-facing readiness", "Medium", "The redesign direction is clear and substantial UI work is present, but mobile polish and final acceptance refinement are still open.", "S03, S09, S13, S16"],
    ["Companion systems ecosystem", "Medium", "Analytics, SDAM, and AVP workstreams are evidenced as real companion systems, but they sit at different maturity and integration stages relative to Tanaghom.", "S17-S19"],
    ["Overall release readiness", "Low to Medium", "The solution is beyond prototype stage, but production identity controls, end-to-end activation, and final acceptance hardening remain open before a client release should be declared.", "S01-S04, S11, S16-S19"],
]

CAPABILITY_ROWS = [
    ["Methodology and planning foundation", "Versioned methodology records, platform cataloging, content-format cataloging, and a 28-day planning model are present in the current application build.", "Implemented", "S05-S06, S10"],
    ["Topic generation and rework cycle", "Topics can be generated, reviewed, sent back for changes, versioned, and regenerated with recorded rationale and reviewer context.", "Implemented", "S07-S09, S15"],
    ["Script generation and quality controls", "Scripts are generated with guardrails, revision handling, anchor checks, and explicit approval checkpoints before downstream progression.", "Implemented", "S07-S08, S14-S15"],
    ["Approval routing and audit trail", "The application contains a real approval engine with multi-stage reviews, audit events, request-change paths, reversible reject behavior, and approval-policy data structures.", "Implemented", "S05, S08, S15"],
    ["Production handoff spine", "The current build contains directive handoff, asset records, and manual production, edit, and distribution stages so the operating model does not stop at script approval.", "Implemented", "S05, S08, S15"],
    ["Operational user experience", "The review surface, workflow views, reviewer context, activity surfaces, and assistant endpoint are already implemented in the inspected build.", "Implemented", "S08-S09, S15-S16"],
    ["Workflow administration", "Workflow versions, draft activation flows, and stage-transition administration exist, but the full client-safe authoring and governance experience is still incomplete.", "Partially Implemented", "S05, S08-S09, S12, S15"],
    ["Methodology administration", "Methodology and content-format data are visible in admin surfaces, but end-user authoring and promotion controls are not complete yet.", "Partially Implemented", "S05, S09-S10, S12"],
    ["Analytics Engine companion system", "The inspected companion system contains official Instagram and YouTube connector components, a normalized analytics read layer, and recent content-oriented dashboards.", "Verified in Companion System", "S17"],
    ["Smart Digital Asset Management companion system", "The inspected companion system contains a ResourceSpace-based asset-library stack with startup automation, proxying, storage-path contracts, and documented role boundaries.", "Verified in Companion System", "S18"],
    ["Agentic Video Producer companion system", "The inspected companion system contains an editor shell, agent-control server, and generation-adapter work. The same evidence also shows remaining stabilization before it should be presented as fully release-ready.", "Verified in Companion System", "S19"],
    ["Security and release controls", "Role and reviewer concepts exist, but full production identity, hardened database authentication, and protected operational controls remain incomplete.", "Partially Implemented", "S03-S05, S08, S12"],
]

DEMO_FINDINGS = [
    ["The review experience is being actively redesigned", "Continuity notes record that the previous client-facing review experience did not meet expectations and that the Lunaris redesign path became the active acceptance track.", "S03"],
    ["The target publishing model has been corrected", "The inspected materials clearly show that Instagram is the publishing target, while Telegram is treated as a control channel rather than a client-facing publishing channel.", "S03, S13"],
    ["Analytics presentation is intentionally honest", "The inspected design guidance explicitly forbids invented analytics and expects empty-state or connection-state handling where live integrations are absent.", "S03, S09"],
    ["The scheduling model already reflects the client's operating rhythm", "A client-origin schedule reference has been brought into the project materials and is aligned with the planning and future calendar experience.", "S06, S13"],
    ["A real application rebuild issue was found and resolved", "Execution tracking and live audit artifacts show that a stale asset issue in the dashboard build was identified and corrected by rebuild and restart.", "S01, S16"],
    ["Responsive refinement is still required", "The latest UI audit still records overflow issues on key operational screens, especially around the home view and workflow administration surfaces.", "S16"],
]

UNVERIFIED_DEMO_NOTES = [
    "Formal meeting minutes, attendee records, and approved client sign-off notes were not found in the inspected materials.",
    "Any verbal or offline demonstration feedback that was not captured in the inspected materials is therefore outside the evidentiary basis of this report.",
]

SOLUTION_CONTEXT_ROWS = [
    [
        "Tanaghom application",
        "Planning, generation, gated approval, production handoff, and operational UI are all present in the current inspected build.",
        "Advanced pilot",
        "This is the primary v1.0 control-plane candidate, but it still needs release hardening and final acceptance refinement.",
    ],
    [
        "Analytics Engine",
        "A companion analytics system was verified with official Instagram and YouTube connector components, a normalized read layer, and three recent content-focused dashboards dated May 25 and June 30, 2026.",
        "Verified companion system",
        "This strengthens overall solution readiness. Official API connectivity is evidenced through temporary or alternate accounts, while the official Moataz account path is intentionally deferred at this stage and should not be read as a failed implementation.",
    ],
    [
        "Smart Digital Asset Management",
        "A companion ResourceSpace-based SDAM stack was verified with Docker orchestration, startup automation, reverse proxying, and documented media-path contracts.",
        "Deployable companion system",
        "This supports the asset-library direction, while live population and production operating proof still need to be demonstrated separately.",
    ],
    [
        "Agentic Video Producer",
        "A companion AVP build was verified with editor-shell, agent-control, and generation-adapter components, together with documented delivery milestones.",
        "Parallel build in stabilization",
        "This is a credible adjacent workstream, but the current evidence still shows remaining stabilization before it should be presented as fully release-ready.",
    ],
]

CHANGE_REQUESTS = [
    {
        "id": "CR-01",
        "title": "Approval Governance And Reviewer-Specific Inbox",
        "description": "Move from stage-only review visibility to user-specific approval visibility with clear assignment context, matched routes, and required approver display.",
        "rationale": "Client and governance users need each reviewer to understand exactly what is awaiting them and why they are allowed to act.",
        "requirements": "User/role/group assignment model; pending approval query; approval-context panel; DB-backed approval policies; auditable assignment snapshots.",
        "technical_impact": "Touches schema, gate engine queries, approval policy reads, reviewer identity plumbing, and operational UI context panels.",
        "repo_impact": "Implemented through principal/role/group tables, approval policy tables, pending-approval queries, API endpoints, and overview/inbox UI additions.",
        "dependencies": "Auth/session source selection, frozen approval semantics, and production authorization controls.",
        "priority": "High",
        "status": "Partially Implemented",
        "gap": "Core data model and UI visibility are present, but the repository still lacks a full production auth source and final semantic closure for all CR01 approval behaviors.",
        "approach": "Convert the current reviewer-selection model into authenticated user identity, keep role/group authorization tests, and close the remaining approval semantics backlog before release.",
        "evidence": "S05, S08-S09, S12, S15",
    },
    {
        "id": "CR-02",
        "title": "Workflow Versioning And Admin Control Plane",
        "description": "Backend and admin UI for workflow versions, stage definitions, transitions, draft editing, and activation.",
        "rationale": "Workflow changes need governed release control rather than YAML-only edits in code branches.",
        "requirements": "Workflow catalog; version drafts; stage/transition editing; activation and rollback-safe promotion path; admin authorization.",
        "technical_impact": "Adds workflow tables, seed/read models, admin APIs, and dashboard admin components.",
        "repo_impact": "Workflow version schema and admin shell are present; self-tests exercise create, update, and activate flows.",
        "dependencies": "Trusted admin identity, release governance, and decisions on whether multi-workflow support is needed now.",
        "priority": "High",
        "status": "Partially Implemented",
        "gap": "The current admin UI is a constrained shell around the existing stage library and does not yet represent a full client-safe workflow authoring product.",
        "approach": "Keep the supported-stage constraint for v1, add authenticated admin access, comparison/rollback UX, and clear release notes per version.",
        "evidence": "S05, S08-S09, S12, S15",
    },
    {
        "id": "CR-03",
        "title": "Methodology And Content-Format Governance",
        "description": "Versioned methodology and content-format management instead of markdown-only control.",
        "rationale": "The client expects methodology and format rules to be governable, inspectable, and eventually editable without ad hoc file surgery.",
        "requirements": "Methodology catalog and versions; content-format registry; platform registry; source digests; admin visibility; future edit/promote flow.",
        "technical_impact": "Adds new schema, loader behavior, read models, and admin UI surfaces while preserving runtime compatibility with the stable format tables.",
        "repo_impact": "Versioned methodology/platform/content-format tables and a read-only admin screen are implemented and seeded from repo source files.",
        "dependencies": "Draft/edit/promote workflow, comparison views, and acceptance rules for methodology changes.",
        "priority": "High",
        "status": "Partially Implemented",
        "gap": "The repo currently proves seeded read models and visibility, but not an end-user authoring workflow or release governance around methodology changes.",
        "approach": "Deliver draft/edit/promote flows next, with checksum comparison and explicit activation history rather than bypassing through direct DB edits.",
        "evidence": "S05, S09-S10, S12, S15",
    },
    {
        "id": "CR-04",
        "title": "Schedule And Calendar Operations Alignment",
        "description": "Align the UI and planning model to the client’s 28-day content schedule reference and operational review habits.",
        "rationale": "The client already manages schedule thinking through a week-grouped, two-slot-per-day artifact; the product needs to mirror that model.",
        "requirements": "Week-grouped calendar/board view; explicit 09:00 and 20:00 slots; post code plus format visibility; override-history preservation; clear template-vs-generated state.",
        "technical_impact": "Uses the existing slot model but requires new read models, UI lenses, and override/audit representation.",
        "repo_impact": "Parametric planning exists and the client schedule reference is imported, but the operational calendar view and override preservation are not yet implemented.",
        "dependencies": "UI lens work, schedule state design, and a decision on where override history should persist.",
        "priority": "Medium-High",
        "status": "Partially Implemented",
        "gap": "The repo supports planning mechanics but not the full schedule-review experience documented in the imported client reference.",
        "approach": "Treat the existing slot model as the base, then add a calendar lens and explicit override audit instead of replacing the planner.",
        "evidence": "S04, S06, S09, S13",
    },
    {
        "id": "CR-05",
        "title": "Operational UX Hardening And Mobile Readiness",
        "description": "Complete the client-facing operational redesign and remove remaining responsive defects.",
        "rationale": "The client-facing review surface must be acceptable on both desktop and mobile, and the redesign is part of restoring stakeholder confidence.",
        "requirements": "Responsive shell, reliable asset delivery, corrected platform mapping, clear overview/workflow views, and no horizontal overflow on key operations pages.",
        "technical_impact": "Frontend shell, state loading, CSS asset handling, responsive layout fixes, and regression coverage.",
        "repo_impact": "The repo contains Lunaris shell and lenses plus resolved stale-asset notes, but the mobile audit still shows overflow defects on key screens.",
        "dependencies": "Completion of the accepted redesign slice and responsive regression testing.",
        "priority": "High",
        "status": "Partially Implemented",
        "gap": "The redesign foundation exists, but the repository’s own audit artifacts show remaining layout defects and unfinished acceptance work.",
        "approach": "Resolve the audited overflow defects first, then complete the planned overview/workflow/calendar acceptance tranche before broader UI ambition.",
        "evidence": "S03, S09, S13, S16",
    },
    {
        "id": "CR-06",
        "title": "External Execution Integrations",
        "description": "Activate real AVP, Postiz, and analytics providers behind the repository’s directive and asset contracts.",
        "rationale": "Manual lifecycle stages are useful for validation, but the client will expect the operating model to connect to downstream execution systems.",
        "requirements": "Stage executors, provider configuration, asset handoff, publish/schedule callback handling, and analytics feedback ingestion.",
        "technical_impact": "Implements adapters against the existing integration contracts and manual-stage data spine.",
        "repo_impact": "The Tanaghom application currently contains the contract declarations and disabled stubs, while companion-system evidence confirms that several of the target external systems already exist outside Tanaghom.",
        "dependencies": "Security hardening, provider access, and prioritized sequence between media-edit, distribution, and analytics.",
        "priority": "High",
        "status": "Partially Implemented",
        "gap": "The connection points inside Tanaghom exist and the companion systems themselves are real, but end-to-end activation between them has not yet been completed in the current Tanaghom build.",
        "approach": "Keep the existing contract boundary, activate one external connection at a time, and preserve the current manual fallback until each activation path is proven operational.",
        "evidence": "S11, S15, S17-S19",
    },
    {
        "id": "CR-07",
        "title": "Production Authentication And IAM Hardening",
        "description": "Replace local-trust and reviewer-selection patterns with production authentication, authorization, and guarded operational controls.",
        "rationale": "Client release requires attributable actions, protected destructive operations, and non-demo-grade security defaults.",
        "requirements": "Session-backed acting user, protected generation/reset endpoints, hardened DB auth, tenant-aware controls, and release-safe admin roles.",
        "technical_impact": "Touches identity plumbing, API authorization, infrastructure config, and deployment runbooks.",
        "repo_impact": "Schema foundations and a trusted-principal proxy exist, but BUILD_STATE explicitly records go-live/IAM hardening as deferred.",
        "dependencies": "Chosen auth provider, deployment model, and final actor/role policy decisions.",
        "priority": "Critical",
        "status": "Partially Implemented",
        "gap": "Foundational role/group data exists, but the repository still reflects local-trust assumptions that are not acceptable for production release.",
        "approach": "Choose the auth source first, remove reviewer-cookie reliance for protected actions, enforce real sessions, and harden infrastructure before any client launch milestone.",
        "evidence": "S03-S05, S08, S12",
    },
    {
        "id": "CR-08",
        "title": "Storyboard And Media-Planning Extension Points",
        "description": "Preserve the future ability to carry storyboard-oriented metadata and UI affordances without reworking the control plane later.",
        "rationale": "The CR01 backlog explicitly reserves storyboard-ready extension points; deferring blindly would risk another schema/UI retrofit later.",
        "requirements": "Storyboard-capable schema fields and a placeholder UI surface or toggle, aligned to future media workflows.",
        "technical_impact": "Minor schema/UI extension if reserved now; larger retrofits if ignored until after release.",
        "repo_impact": "The need is recorded in the CR01 backlog, but no implementation was found in the repository.",
        "dependencies": "Core v1.0 scope discipline and confirmation that storyboard support is still desired in the near term.",
        "priority": "Medium",
        "status": "Not Yet Implemented",
        "gap": "The repository contains only backlog intent, not code or schema delivery.",
        "approach": "Keep as a thin extension reservation in the post-v1.0 sequence unless immediate business need elevates it.",
        "evidence": "S12",
    },
]

ALIGNMENT_ROWS = [
    [cr["id"], cr["title"], cr["status"], cr["evidence"], cr["gap"]] for cr in CHANGE_REQUESTS
]

DASHBOARD_ROWS = [
    ["Tanaghom planning and run creation", "5/5", "Implemented and verified", "Planning logic and run creation are live in the current application build and passed automated verification."],
    ["Tanaghom topic and script generation", "4/5", "Implemented", "Generation and revision loops are real, with further quality tuning still appropriate before release."],
    ["Tanaghom approval engine and audit", "5/5", "Implemented and verified", "The gate lifecycle, sign-off behavior, and audit trail are real and test-backed."],
    ["Tanaghom production handoff spine", "4/5", "Implemented", "Manual production and distribution stages are modeled and connected through directives and asset records."],
    ["Tanaghom workflow administration", "3/5", "Partially implemented", "Workflow versioning exists, but client-safe governance and full administration completion remain open."],
    ["Tanaghom methodology administration", "3/5", "Partially implemented", "Catalog visibility is available, while business-authoring and promotion are still open."],
    ["Tanaghom client-facing UX readiness", "2/5", "Partially implemented", "The redesign direction is strong, but responsive refinement and final acceptance polish remain necessary."],
    ["Analytics Engine companion system", "4/5", "Verified in companion system", "Official-connector components and recent content dashboards are present; the official Moataz account path is intentionally deferred."],
    ["SDAM companion system", "3/5", "Verified in companion system", "A deployable ResourceSpace-based stack exists, with final operational population still to be demonstrated."],
    ["AVP companion system", "3/5", "Verified in companion system", "Core platform and initial interface work are real, with stabilization still in progress."],
    ["End-to-end cross-system activation", "2/5", "Partially implemented", "Tanaghom-side contracts exist, but not all companion capabilities are yet activated through Tanaghom."],
    ["Production identity and security controls", "2/5", "Partially implemented", "Foundations exist, but production-grade identity and release controls are not yet complete."],
]

ROADMAP_ROWS = [
    ["Release 0.91", "Client-facing stabilization baseline", "Complete the current UI stabilization pass, close the audited responsive issues, preserve a clean application build, and issue the enriched readiness pack as the working baseline.", "Relies on the current Tanaghom build and audit evidence only.", "Creates a trustworthy baseline for all subsequent release discussions."],
    ["Release 0.95", "Governance and admin completion", "Finish approval-governance completion, workflow-version administration, methodology and content-format governance, and authenticated reviewer/admin identity.", "Depends on approval semantics closure and identity-source decisions.", "Allows the client to evaluate a governed operating platform rather than a promising but incomplete control layer."],
    ["Release 0.98", "Operational acceptance slice", "Deliver the client-aligned schedule and calendar experience, preserve override history, and finish the acceptance-grade review experience across desktop and mobile.", "Depends on the governance layer and the ongoing Lunaris redesign track.", "Moves the platform from functional depth to day-to-day operational usability."],
    ["Release 0.99", "Companion-system activation", "Connect the prioritized external systems into Tanaghom in sequence: analytics feedback, production/distribution activation, and asset-library operating alignment.", "Depends on security hardening, credential availability, and final sequencing decisions across Tanaghom, Analytics, SDAM, and AVP.", "Converts separate system maturity into an integrated operating model."],
    ["Release 1.0", "Controlled client release", "Run a clean pilot cycle, complete UAT, finalize runbooks and support boundaries, and release from a frozen and governed baseline.", "Depends on successful completion of 0.99 and explicit stakeholder sign-off.", "Positions v1.0 as a controlled release baseline with clear governance, rather than an advanced pilot presented too early."],
]

RISK_ROWS = [
    ["Production identity and security controls are not yet complete", "Critical", "Execution tracking explicitly records deferred go-live hardening and local-trust assumptions in the present environment.", "Complete identity, authorization, endpoint protection, and database hardening before any client release milestone."],
    ["The currently inspected build is ahead of a formal release package", "High", "The active build contains substantial delivery work that is real, but still needs formal release packaging and release governance.", "Freeze an explicit release baseline before any go-live commitment is made."],
    ["User experience polish remains visible to stakeholders", "High", "The latest UI audit still records responsive defects on key operational surfaces.", "Treat responsive hardening as an immediate release-readiness task rather than a cosmetic follow-up."],
    ["Cross-system activation is not yet complete", "High", "Tanaghom contains connection seams, and companion systems exist, but the systems are not yet fully activated together end to end.", "Sequence external-system activation deliberately and keep current manual fallbacks until each path is proven."],
    ["Generated-content quality still needs final operational tuning", "Medium", "The quality backlog records open dialect, tone, and metaphor issues that were acceptable for flow validation but not ideal for final release confidence.", "Run a targeted quality pass and keep mandatory human review on sensitive content paths."],
    ["Demonstration data can be affected by verification activity", "Medium", "The automated verification suites create and update test rounds and workflow versions in the active environment.", "Use a controlled demo dataset and a cleanup discipline before formal client walkthroughs."],
]

RECOMMENDATION_ROWS = [
    ["Release framing", "Present the current position as a strong advanced pilot moving toward v1.0, not as a fully completed end-state."],
    ["Control-plane completion", "Finish the approval, workflow, identity, and methodology-governance items before adding broader surface complexity."],
    ["Client-facing UX sequence", "Resolve the audited responsive defects and complete the accepted review/schedule experience before broader visual ambition."],
    ["External activation strategy", "Connect Analytics, SDAM, AVP, and distribution systems through the existing Tanaghom seams in a controlled sequence instead of bypassing the operating model."],
    ["Acceptance management", "Use a curated demo dataset, explicit release baseline, and formal acceptance checklist before the client UAT cycle."],
]

ACCEPTANCE_ROWS = [
    ["Tanaghom planning, generation, and approval core", "Yes", "Acceptable for a controlled pilot because the main editorial lifecycle is real and automated verification supports it."],
    ["Tanaghom operational UI", "Conditional", "Accept after the current responsive issues are closed and the agreed acceptance slice of the redesign is completed."],
    ["Workflow and methodology administration", "Conditional", "Accept after authenticated administration and clearer promotion/release governance are completed."],
    ["Analytics companion system", "Conditional", "Accept as proven companion-system progress now; accept as part of Tanaghom v1.0 only after the agreed activation path is connected."],
    ["SDAM companion system", "Conditional", "Accept as infrastructure readiness now; accept as part of v1.0 after live operational setup and role usage are proven."],
    ["AVP companion system", "Conditional", "Accept as parallel delivery progress now; accept as part of v1.0 only after remaining stabilization is closed."],
    ["Production identity and security controls", "No", "Not acceptable for release until the current local-trust assumptions are replaced by production-grade controls."],
]

COMPONENT_ROWS = [
    ["Planner", "planner/plan_round.py; planner/README.md", "Parametric round creation, ratio scaling, cursor/lens rules."],
    ["Writers", "agents/run_writers.py; agents/providers.py", "Topic/script generation, rework, quality guards, fallback providers."],
    ["Gate engine", "gates/engine.py", "Stage state machine, approvals, directives, DAM, actor model seams."],
    ["Gate API", "gates/api.py", "HTTP surface for rounds, jobs, approvals, admin data, assets, and assistant."],
    ["Methodology loader", "loader/load_methodology.py", "Imports markdown canon and HCS into runtime and versioned tables."],
    ["Operational dashboard", "dashboard/lib/review-context.tsx; dashboard/components/review/*", "Review queue, overview, workflow lens, stage actions, assistant panel."],
    ["Workflow admin", "dashboard/components/admin/workflow-admin.tsx", "Versioned workflow shell and policy editor."],
    ["Methodology admin", "dashboard/components/admin/methodology-admin.tsx", "Read-only methodology/format/platform visibility."],
    ["Integration seams", "integrations/contracts.py; integrations/stubs.py", "Declared AVP/Postiz/analytics stage executor contracts."],
    ["Verification suites", "gates/selftest.py; gates/api_selftest.py; gates/lifecycle_selftest.py; dashboard/e2e/*", "Repository verification coverage."],
    ["Analytics companion system", "Brand Shield Repo", "External analytics engine evidence reviewed for overall solution readiness context."],
    ["SDAM companion system", "resourcespace-stitch-stack", "External asset-library stack evidence reviewed for overall solution readiness context."],
    ["AVP companion system", "Agentic-Video-Producer", "External video-production workstream reviewed for overall solution readiness context."],
]

VERIFICATION_ROWS = [
    ["`curl http://localhost:8009/health`", "Passed", "Gate API returned HTTP 200 with `{\\\"ok\\\": true}` during this audit."],
    ["`curl -I http://localhost:3000`", "Passed", "Dashboard returned HTTP 200 during this audit."],
    ["`docker exec ... python -m gates.selftest`", "Passed", "Engine/self-test suite completed successfully on July 2, 2026."],
    ["`docker exec ... python -m gates.api_selftest`", "Passed", "API/self-test suite completed successfully on July 2, 2026."],
    ["`docker exec ... python -m gates.lifecycle_selftest`", "Passed", "Lifecycle self-test suite completed successfully on July 2, 2026."],
]

VISUAL_REFERENCES = [
    {
        "title": "Tanaghom - current review surface",
        "caption": "Current Tanaghom review experience from the inspected application build.",
        "path": ROOT / "docs" / "assets" / "program_evolution_report_v1_0" / "tanaghom-review-surface.png",
        "width": 3.6,
    },
    {
        "title": "Tanaghom - redesign workflow reference",
        "caption": "Workflow reference used in the active redesign track, showing the corrected Instagram publishing target and the broader operating flow.",
        "path": ROOT / "docs" / "assets" / "program_evolution_report_v1_0" / "tanaghom-workflow-reference.png",
        "width": 5.8,
    },
    {
        "title": "Tanaghom - Telegram Agent Interaction Channel",
        "caption": "Telegram agent interaction channel, live and active as a control channel with ongoing UI and UX refinement.",
        "path": ROOT / "docs" / "assets" / "program_evolution_report_v1_0" / "tanaghom-telegram-summary.png",
        "width": 2.1,
    },
    {
        "title": "Tanaghom - Telegram Approval Cards",
        "caption": "Telegram approval-card flow showing actionable review items surfaced in the live control channel.",
        "path": ROOT / "docs" / "assets" / "program_evolution_report_v1_0" / "tanaghom-telegram-approval.png",
        "width": 2.1,
    },
    {
        "title": "Analytics Engine - workflow automation view",
        "caption": "Companion analytics-system workflow view showing sentiment tagging automation in the inspected n8n interface.",
        "path": ROOT / "docs" / "assets" / "program_evolution_report_v1_0" / "analytics-n8n-workflow.png",
        "width": 5.8,
    },
    {
        "title": "Analytics Engine - dashboard views",
        "caption": "Content analytics dashboard view from the companion analytics system. Official Instagram and YouTube API connectivity is evidenced through temporary or alternate accounts while the official Moataz account onboarding remains intentionally deferred for now.",
        "path": ROOT / "docs" / "assets" / "program_evolution_report_v1_0" / "analytics-dashboard-overview.png",
        "width": 6.0,
    },
    {
        "title": "SDAM - asset detail view",
        "caption": "ResourceSpace-based SDAM asset-detail view from the companion asset-library stack.",
        "path": ROOT / "docs" / "assets" / "program_evolution_report_v1_0" / "sdam-asset-detail.png",
        "width": 4.0,
    },
    {
        "title": "SDAM - collection browser",
        "caption": "ResourceSpace-based SDAM collection and browsing view showing assets organized in the companion library stack.",
        "path": ROOT / "docs" / "assets" / "program_evolution_report_v1_0" / "sdam-collection-browser.png",
        "width": 6.0,
    },
    {
        "title": "AVP - live editing surface",
        "caption": "Agentic Video Producer live editing surface and proposal-preview experience from the current parallel workstream.",
        "path": ROOT / "docs" / "assets" / "program_evolution_report_v1_0" / "avp-live-editing-surface.png",
        "width": 6.0,
    },
    {
        "title": "AVP - target concept dual surface",
        "caption": "Agentic Video Producer target concept showing editing and co-creation surfaces in the under-development vision.",
        "path": ROOT / "docs" / "assets" / "program_evolution_report_v1_0" / "avp-target-concept-dual-surface.png",
        "width": 6.0,
    },
    {
        "title": "AVP - target concept media controls",
        "caption": "Agentic Video Producer target concept for media controls and human co-creation in the future-state vision.",
        "path": ROOT / "docs" / "assets" / "program_evolution_report_v1_0" / "avp-target-concept-media-controls.png",
        "width": 6.0,
    },
]



def build_markdown() -> str:
    lines: list[str] = []
    lines.append(f"# {TITLE}")
    lines.append("")
    lines.append(SUBTITLE)
    lines.append("")
    lines.append("## Document Control")
    lines.append("")
    lines.append(f"- Version: {VERSION}")
    lines.append(f"- Date: {REPORT_DATE}")
    lines.append(f"- Classification: {CLASSIFICATION}")
    lines.append(f"- Shared reference baseline: {SHARED_REFERENCE_BASELINE}")
    lines.append(f"- Assessment basis: {ASSESSMENT_BASIS}")
    lines.append(f"- Internal build reference: `{BRANCH}`")
    lines.append("")
    lines.append("## 1 Executive Summary")
    lines.append("")
    lines.append("This report is written as a client-facing program update. It distinguishes clearly between:")
    lines.append("")
    lines.append("- What is already working inside the current Tanaghom application build")
    lines.append("- What is partly completed and still needs delivery work")
    lines.append("- What has been verified in companion systems outside Tanaghom")
    lines.append("- What remains to be completed before a controlled v1.0 release should be presented")
    lines.append("")
    for item in EXEC_SUMMARY:
        lines.append(f"- {item}")
    lines.append("")
    lines.append("## 2 Repository Maturity Assessment")
    lines.append("")
    lines.append("Although this section retains the requested title, the assessment below is written in business terms and refers to the current application build and broader delivery program.")
    lines.append("")
    lines.append("| Area | Maturity | Assessment | Evidence |")
    lines.append("|---|---|---|---|")
    for row in MATURITY_ROWS:
        lines.append("| " + " | ".join(row) + " |")
    lines.append("")
    lines.append("## 3 Current Implemented Capabilities")
    lines.append("")
    lines.append("The table below focuses on capabilities that are evidenced either in the current Tanaghom application or in separately inspected companion systems that the client asked to have reflected in the report.")
    lines.append("")
    lines.append("| Capability | Verified repository status | Status | Evidence |")
    lines.append("|---|---|---|---|")
    for row in CAPABILITY_ROWS:
        lines.append("| " + " | ".join(row) + " |")
    lines.append("")
    lines.append("What this means today:")
    lines.append("")
    lines.append("- Tanaghom already has a real editorial control layer rather than only conceptual planning materials.")
    lines.append("- Companion systems for Analytics, SDAM, and AVP are genuine workstreams with inspected evidence, but they should not yet be described as fully activated through Tanaghom.")
    lines.append("- Client communication should therefore present the solution as substantially advanced, while still being explicit about activation and release work that remains.")
    lines.append("")
    lines.append("## 4 Client Demonstration Summary")
    lines.append("")
    lines.append("The current project materials do not include formal client meeting minutes. This section therefore summarizes only what is evidenced in the inspected materials and the verified companion-system review.")
    lines.append("")
    lines.append("| Recorded finding | Repository evidence | Source |")
    lines.append("|---|---|---|")
    for row in DEMO_FINDINGS:
        lines.append("| " + " | ".join(row) + " |")
    lines.append("")
    lines.append("Unverified from repository:")
    for note in UNVERIFIED_DEMO_NOTES:
        lines.append(f"- {note}")
    lines.append("")
    lines.append("### Overall Solution Landscape Relevant To The Client")
    lines.append("")
    lines.append("| Workstream | What was verified | Current status | Client-facing implication |")
    lines.append("|---|---|---|---|")
    for row in SOLUTION_CONTEXT_ROWS:
        lines.append("| " + " | ".join(row) + " |")
    lines.append("")
    lines.append("Important nuance regarding Moataz onboarding:")
    lines.append("")
    lines.append("- The official Moataz account path for the Analytics Engine is intentionally deferred at this stage. It should therefore be described as a sequencing decision, not as a failed implementation.")
    lines.append("")
    lines.append("## 5 Consolidated Change Requests")
    lines.append("")
    for cr in CHANGE_REQUESTS:
        lines.append(f"### {cr['id']} {cr['title']}")
        lines.append("")
        for label, value in (
            ("Description", cr["description"]),
            ("Business rationale", cr["rationale"]),
            ("Functional requirements", cr["requirements"]),
            ("Technical impact", cr["technical_impact"]),
            ("Current platform impact", cr["repo_impact"]),
            ("Dependencies", cr["dependencies"]),
            ("Priority", cr["priority"]),
            ("Current implementation status", cr["status"]),
            ("Gap analysis", cr["gap"]),
            ("Recommended implementation approach", cr["approach"]),
            ("Evidence", cr["evidence"]),
        ):
            lines.append(f"- **{label}:** {value}")
        lines.append("")
    lines.append("## 6 Repository Alignment Matrix")
    lines.append("")
    lines.append("For client readability, this matrix uses the statuses `Implemented`, `Partially Implemented`, and `Not Yet Implemented` rather than the harsher wording often used in engineering backlogs.")
    lines.append("")
    lines.append("| CR ID | Title | Status | Evidence | Gap summary |")
    lines.append("|---|---|---|---|---|")
    for row in ALIGNMENT_ROWS:
        lines.append("| " + " | ".join(row) + " |")
    lines.append("")
    lines.append("## 7 Progress Dashboard")
    lines.append("")
    lines.append("This dashboard combines the current Tanaghom application with the externally reviewed companion systems so the client can see overall solution readiness at a glance.")
    lines.append("")
    lines.append("| Module | Maturity | Current state | Note |")
    lines.append("|---|---|---|---|")
    for row in DASHBOARD_ROWS:
        lines.append("| " + " | ".join(row) + " |")
    lines.append("")
    lines.append("## 8 Release Roadmap")
    lines.append("")
    lines.append("The sequence below is a recommended release path derived from verified evidence. It is not a claim that all releases are already approved or funded.")
    lines.append("")
    lines.append("| Release | Objective | Included work | Dependencies | Exit rationale |")
    lines.append("|---|---|---|---|---|")
    for row in ROADMAP_ROWS:
        lines.append("| " + " | ".join(row) + " |")
    lines.append("")
    lines.append("## 9 Risks and Recommendations")
    lines.append("")
    lines.append("### Risks")
    lines.append("")
    lines.append("| Risk | Severity | Evidence | Recommendation |")
    lines.append("|---|---|---|---|")
    for row in RISK_ROWS:
        lines.append("| " + " | ".join(row) + " |")
    lines.append("")
    lines.append("### Recommendations")
    lines.append("")
    for row in RECOMMENDATION_ROWS:
        lines.append(f"- **{row[0]}:** {row[1]}")
    lines.append("")
    lines.append("## 10 Client Acceptance Considerations")
    lines.append("")
    lines.append("| Area | Accept now? | Condition |")
    lines.append("|---|---|---|")
    for row in ACCEPTANCE_ROWS:
        lines.append("| " + " | ".join(row) + " |")
    lines.append("")
    lines.append("## 11 Appendix")
    lines.append("")
    lines.append("### Appendix A. Source Register")
    lines.append("")
    lines.append("| Code | Title | Path | Use |")
    lines.append("|---|---|---|---|")
    for s in SOURCE_REFS:
        lines.append(f"| {s.code} | {s.title} | `{s.path}` | {s.use} |")
    lines.append("")
    lines.append("### Appendix B. Affected Repository Components")
    lines.append("")
    lines.append("| Component | Key files | Relevance |")
    lines.append("|---|---|---|")
    for row in COMPONENT_ROWS:
        lines.append("| " + " | ".join(row) + " |")
    lines.append("")
    lines.append("### Appendix C. Verification Performed During This Audit")
    lines.append("")
    lines.append("| Check | Result | Note |")
    lines.append("|---|---|---|")
    for row in VERIFICATION_ROWS:
        lines.append("| " + " | ".join(row) + " |")
    lines.append("")
    lines.append("### Appendix D. Visual References")
    lines.append("")
    for item in VISUAL_REFERENCES:
        lines.append(f"- **{item['title']}:** {item['caption']}")
        if item["path"] is not None:
            lines.append(f"  - Source image: `{item['path']}`")
        else:
            lines.append("  - Source image: Placeholder reserved for later client-supplied capture.")
    lines.append("")
    return "\n".join(lines).strip() + "\n"


def build_docx() -> Document:
    doc = Document()
    style_document(doc)
    enable_update_fields_on_open(doc)
    props = doc.core_properties
    props.title = TITLE
    props.subject = "Tanaghom release readiness and change request integration"
    props.author = "OpenAI Codex"
    props.category = "Client report"
    props.comments = "Evidence-based client-facing assessment"

    section = doc.sections[0]
    add_footer(section)

    if LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(LOGO), width=Inches(1.55))

    add_text(doc, TITLE, size=22, bold=True, color=COLOR_NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, after=6)
    add_text(doc, SUBTITLE, size=13.5, color=COLOR_MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=18)
    add_text(doc, f"Version {VERSION} | {REPORT_DATE}", size=11, color=COLOR_BLUE_DARK,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    add_text(doc, CLASSIFICATION, size=10.5, italic=True, color=COLOR_MUTED,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=18)
    add_label_value_table(doc, [
        ("Prepared for", "Tanaghom AI Content Department stakeholders"),
        ("Document purpose", "Client communication and internal program governance"),
        ("Assessment basis", ASSESSMENT_BASIS),
        ("Internal build reference", BRANCH),
        ("Release framing", "Controlled phased path to v1.0"),
    ])
    doc.add_page_break()

    add_text(doc, "Document Control", style="Heading 1")
    add_label_value_table(doc, [
        ("Document title", TITLE),
        ("Subtitle", "Client Change Request Integration, Implementation Status & Release Roadmap"),
        ("Target release", "v1.0"),
        ("Version", VERSION),
        ("Issue date", REPORT_DATE),
        ("Classification", CLASSIFICATION),
        ("Shared reference baseline", SHARED_REFERENCE_BASELINE),
        ("Assessment scope", "Tanaghom application build, tracked documentation, automated verification, live health checks, and selected companion-system evidence"),
    ])

    add_text(doc, "Revision History", style="Heading 2")
    add_matrix_table(
        doc,
        ["Version", "Date", "Summary", "Author"],
        [
            ["1.0", REPORT_DATE, "Initial repository evidence audit and canonical release-readiness issue.", "OpenAI Codex"],
            ["1.1", REPORT_DATE, "Client-facing wording refined and external parallel workstreams recorded with explicit repository-verification boundaries.", "OpenAI Codex"],
            ["1.2", REPORT_DATE, "Expanded client-facing narrative, verified companion-system status, and visual appendix added.", "OpenAI Codex"],
        ],
        [0.75, 1.05, 4.0, 0.7],
    )

    add_text(doc, "Contents", style="Heading 1")
    add_text(doc, "This Word document includes an automatic table of contents field. If page numbers are not shown immediately, update fields once in Microsoft Word.", size=10.5, color=COLOR_MUTED, italic=True, after=8)
    toc_paragraph = doc.add_paragraph()
    toc_paragraph.paragraph_format.space_after = Pt(8)
    add_toc_field(toc_paragraph)
    doc.add_page_break()

    add_text(doc, "1 Executive Summary", style="Heading 1")
    add_text(doc, "This report is intended to support both client communication and internal governance. It separates verified current capability, verified companion-system progress, and remaining delivery work before a controlled v1.0 release should be presented.", after=8)
    add_text(doc, "How to read this report", style="Heading 2")
    add_bullet(doc, "Implemented means the capability is evidenced in the inspected Tanaghom build or in a separately inspected companion system that is clearly identified as such.")
    add_bullet(doc, "Partially Implemented means some meaningful delivery work is present, but the capability is not yet complete or not yet activated end to end.")
    add_bullet(doc, "Not Yet Implemented is used where the requested client capability is still planned rather than delivered.")
    for item in EXEC_SUMMARY:
        add_bullet(doc, item)

    add_text(doc, "2 Repository Maturity Assessment", style="Heading 1")
    add_text(doc, "Although this section retains the requested title, the assessment below is written in business terms and refers to the current application build and broader delivery program rather than engineering-only language.", after=8)
    add_matrix_table(doc, ["Area", "Maturity", "Assessment", "Evidence"], MATURITY_ROWS, [1.5, 1.1, 2.75, 1.0])

    add_text(doc, "3 Current Implemented Capabilities", style="Heading 1")
    add_text(doc, "The table below combines capabilities evidenced in Tanaghom itself with companion-system capabilities that the client requested to be represented for overall solution-readiness visibility.", after=8)
    add_matrix_table(doc, ["Capability", "Verified current status", "Status", "Evidence"], CAPABILITY_ROWS, [1.45, 3.15, 1.1, 0.8], status_col=2)
    add_text(doc, "Client Interpretation", style="Heading 2")
    add_bullet(doc, "Tanaghom is already a real editorial operating platform rather than only a planning concept.")
    add_bullet(doc, "Analytics, SDAM, and AVP are genuine parallel workstreams with verified evidence, but they are not yet all activated end to end through Tanaghom.")
    add_bullet(doc, "The strongest client position today is therefore substantial overall readiness with staged release sequencing still required.")

    add_text(doc, "4 Client Demonstration Summary", style="Heading 1")
    add_text(doc, "The inspected materials do not include formal client meeting minutes. The summary below therefore reflects only what is evidenced in the current project materials and the verified companion-system review.", after=8)
    add_matrix_table(doc, ["Recorded finding", "Evidence summary", "Source"], DEMO_FINDINGS, [1.65, 3.95, 0.85])
    add_text(doc, "Unable To Verify From Reviewed Materials", style="Heading 2")
    for note in UNVERIFIED_DEMO_NOTES:
        add_bullet(doc, note)
    add_text(doc, "Overall Solution Landscape Relevant To The Client", style="Heading 2")
    add_text(doc, "The client asked for the external Analytics, SDAM, and AVP workstreams to be reflected in the document. The table below records only what was verified in the separately inspected companion-system materials.", after=8)
    add_matrix_table(doc, ["Workstream", "What was verified", "Current status", "Client-facing implication"], SOLUTION_CONTEXT_ROWS, [1.15, 2.5, 1.2, 1.55])
    add_text(doc, "Moataz Analytics Onboarding Nuance", style="Heading 3")
    add_text(doc, "The official Moataz account path for the Analytics Engine is intentionally deferred at this stage. In client terms, that should be described as a conscious sequencing decision rather than as a failed implementation.", after=8)

    add_text(doc, "5 Consolidated Change Requests", style="Heading 1")
    add_text(doc, "The consolidated change requests below bring together the recorded CR01 backlog, demonstrated operating needs, release-governance gaps, and the activation work needed to convert separate system progress into a coherent client release.", after=8)
    for cr in CHANGE_REQUESTS:
        add_text(doc, f"{cr['id']} {cr['title']}", style="Heading 2")
        add_label_value_table(doc, [
            ("Description", cr["description"]),
            ("Business rationale", cr["rationale"]),
            ("Functional requirements", cr["requirements"]),
            ("Technical impact", cr["technical_impact"]),
            ("Current platform impact", cr["repo_impact"]),
            ("Dependencies", cr["dependencies"]),
            ("Priority", cr["priority"]),
            ("Current implementation status", cr["status"]),
            ("Gap analysis", cr["gap"]),
            ("Recommended implementation approach", cr["approach"]),
            ("Evidence", cr["evidence"]),
        ])

    add_text(doc, "6 Repository Alignment Matrix", style="Heading 1")
    add_text(doc, "For client readability, this matrix uses `Not Yet Implemented` instead of harsher engineering backlog language. The classifications remain evidence-based.", after=8)
    add_matrix_table(doc, ["CR ID", "Title", "Status", "Evidence", "Gap summary"], ALIGNMENT_ROWS, [0.75, 1.8, 1.15, 0.9, 1.85], status_col=2)

    add_text(doc, "7 Progress Dashboard", style="Heading 1")
    add_text(doc, "Maturity scale used in this dashboard: 5 = implemented and operationally verified, 3 = materially present but not complete, 1 = concept or backlog only. This dashboard includes the verified companion systems so the client can see the broader solution picture.", after=8)
    add_matrix_table(doc, ["Module", "Maturity", "Current state", "Note"], DASHBOARD_ROWS, [1.95, 0.7, 1.4, 2.3], status_col=2)

    add_text(doc, "8 Release Roadmap", style="Heading 1")
    add_text(doc, "The roadmap below is a recommended sequence derived from verified evidence and open delivery gaps. It should be read as a practical release path rather than a claim that all release steps are already approved.", after=8)
    add_matrix_table(doc, ["Release", "Objective", "Included work", "Dependencies", "Exit rationale"], ROADMAP_ROWS, [0.95, 1.3, 2.25, 1.2, 0.9])
    add_text(doc, "Roadmap Rationale", style="Heading 2")
    add_bullet(doc, "Tanaghom governance and user experience should be stabilized before cross-system activation is attempted at scale.")
    add_bullet(doc, "Companion-system maturity is valuable immediately for client confidence, but companion-system existence is not the same as end-to-end activation inside Tanaghom.")
    add_bullet(doc, "The intentionally deferred official Moataz analytics path should be activated only when the surrounding governance and release controls are ready to support it.")

    add_text(doc, "9 Risks and Recommendations", style="Heading 1")
    add_text(doc, "Key Risks", style="Heading 2")
    add_matrix_table(doc, ["Risk", "Severity", "Evidence", "Recommendation"], RISK_ROWS, [1.75, 0.8, 1.8, 2.15])
    add_text(doc, "Implementation Recommendations", style="Heading 2")
    add_matrix_table(doc, ["Recommendation Area", "Recommended action"], RECOMMENDATION_ROWS, [1.7, 4.8])

    add_text(doc, "10 Client Acceptance Considerations", style="Heading 1")
    add_text(doc, "The acceptance view below separates what can reasonably be accepted now for a controlled pilot from what should only be accepted once release-grade completion and activation are achieved.", after=8)
    add_matrix_table(doc, ["Area", "Accept now?", "Condition"], ACCEPTANCE_ROWS, [2.1, 1.0, 3.4], status_col=1)

    add_text(doc, "11 Appendix", style="Heading 1")
    add_text(doc, "Appendix A. Source Register", style="Heading 2")
    add_matrix_table(
        doc,
        ["Code", "Title", "Path", "Use"],
        [[s.code, s.title, s.path, s.use] for s in SOURCE_REFS],
        [0.55, 1.55, 1.9, 2.5],
    )
    add_text(doc, "Appendix B. Affected Repository Components", style="Heading 2")
    add_matrix_table(doc, ["Component", "Key files", "Relevance"], COMPONENT_ROWS, [1.35, 2.35, 2.8])
    add_text(doc, "Appendix C. Verification Performed During This Audit", style="Heading 2")
    add_matrix_table(doc, ["Check", "Result", "Note"], VERIFICATION_ROWS, [2.3, 0.8, 3.4], status_col=1)
    add_text(doc, "Appendix D. Visual References", style="Heading 2")
    add_text(doc, "The client asked for screenshots where possible. The visual appendix below therefore mixes verified current screenshots with clearly labeled placeholders reserved for later insertion.", after=8)
    for item in VISUAL_REFERENCES:
        add_text(doc, item["title"], style="Heading 3")
        if item["path"] is not None and Path(item["path"]).exists():
            add_visual(doc, Path(item["path"]), item["caption"], width=item["width"])
        else:
            add_placeholder(doc, item["title"], item["caption"])
    return doc


def main() -> None:
    ensure_dirs()
    MD_PATH.write_text(build_markdown(), encoding="utf-8")
    doc = build_docx()
    doc.save(DOCX_PATH)
    print(f"wrote {MD_PATH}")
    print(f"wrote {DOCX_PATH}")


if __name__ == "__main__":
    main()
